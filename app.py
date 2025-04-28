import streamlit as st
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.pydantic_v1 import Field
from pydantic import BaseModel
from langgraph.graph import StateGraph
from typing import List, Annotated, Literal, Sequence, TypedDict
from langgraph.graph import END, StateGraph, START
import asyncio
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.schema import Document
from dotenv import load_dotenv
from langsmith import Client
import streamlit_authenticator as stauth
import pandas as pd
import matplotlib.pyplot as plt
import networkx as nx
import japanize_matplotlib 
import numpy as np
import seaborn as sns
import json
from st_login_form import login_form
from datetime import datetime
from janome.tokenizer import Tokenizer
from wordcloud import WordCloud
from google.oauth2 import service_account
import uuid
import time
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.errors import HttpError
import openai
from langchain.chat_models import ChatOpenAI
import re
from docx import Document
from langchain.prompts import ChatPromptTemplate
from langchain.chains import LLMChain
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.preprocessing import MultiLabelBinarizer
import plotly.graph_objects as go
from google.cloud import vision
from pdf2image import convert_from_path
import tempfile
from io import BytesIO
from PIL import Image
import cv2
from langdetect import detect
import openpyxl
import io
from langchain.vectorstores import FAISS
from langchain.schema.document import Document as LangChainDocument




# フォントパスの指定（必要に応じて利用）
font_path1 = "./font/NotoSansJP-Regular.ttf"

# .envファイルを読み込む
load_dotenv(dotenv_path=".env.example")

# JSON文字列を使う場合（注意：evalやjson.loadsが必要）

SERVICE_ACCOUNT_INFO = st.secrets["GOOGLE_DRIVE_CREDENTIAL_JSON"]
VISION_ACCOUNT_INFO = st.secrets["GOOGLE_APPLICATION_CREDENTIALS"]

# Google API認証（JSON文字列の場合）
creds = service_account.Credentials.from_service_account_info(
    SERVICE_ACCOUNT_INFO,
    scopes=['https://www.googleapis.com/auth/drive']
)
drive_service = build('drive', 'v3', credentials=creds)

# Vision API クライアント設定
oS_creds = service_account.Credentials.from_service_account_info(
    VISION_ACCOUNT_INFO
)
vision_client = vision.ImageAnnotatorClient(credentials=oS_creds)

# パスワードを環境変数から取得
PASSWORD = os.getenv("PASSWORD")
PDF_input_folder = os.getenv("PDF_input_folder")
DOCS_output_folder = os.getenv("DOCS_output_folder")
GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
#USER_AGENT = os.getenv("USER_AGENT")
#os.environ['USER_AGENT'] = USER_AGENT

# パスワード入力ボックス
inputText_A = st.text_input('パスワードを入力してください', type="password")

# 環境変数`PASSWORD`と比較
if inputText_A == PASSWORD:
        
    # ---- ログイン成功後に表示するアプリのメイン機能をここに記述 ----
    st.write("ここにチャットシステムやデータ分析モードなどを切り替えるコードを実装できます。")

    # --- モード選択 ---
    mode = st.sidebar.radio("モード選択", ( "特許要約システム", "データ解析システム", "対比表作成システム"))

    if mode == "特許要約システム":
        # --- Google Drive 連携関数 ---
        st.title("📄特許要約システム")
        st.write("特許PDFファイル（現時点では日本語特許のみ）を読み込み、文献の中身をChatGPTが確認し、表にまとめます！一時的に私用のgoogledriveに保管させて頂きますので、機密情報は絶対にアップロードしないでください。まだ要約イマイチかもしれないです。")
        st.image("./fig/特許要約システム1.jpg", use_container_width=True)
        st.image("./fig/特許要約システム2.jpg", use_container_width=True)
        
        
        def upload_to_drive(folder_id, file_path, mime_type='application/pdf'):
            unique_name = f"{os.path.splitext(os.path.basename(file_path))[0]}_{int(time.time())}_{uuid.uuid4().hex[:6]}.pdf"
            media = MediaFileUpload(file_path, mimetype=mime_type)
            file_metadata = {'name': unique_name, 'parents': [folder_id]}
            uploaded = drive_service.files().create(body=file_metadata, media_body=media, fields='id,name').execute()
            return uploaded['id'], uploaded['name']

        def delete_from_drive(file_id):
            try:
                drive_service.files().delete(fileId=file_id).execute()
            except HttpError as e:
                print(f"削除失敗: {e}")  # ログだけ残す

        def convert_pdf_to_doc(file_id, output_folder_id):
            copied = drive_service.files().copy(fileId=file_id, body={
                'mimeType': 'application/vnd.google-apps.document',
                'parents': [output_folder_id]
            }).execute()
            return copied['id']

        def download_doc_as_docx(file_id, local_path):
            try:
                request = drive_service.files().export_media(fileId=file_id,
                                                            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                with open(local_path, 'wb') as f:
                    f.write(request.execute())
                return local_path
            except HttpError as error:
                print(f"❌ ドキュメントのダウンロード失敗: {error}")
                return None
            
        def preprocess_image_for_ocr(pil_image):
            """OCR前の画像前処理（OpenCV使用）"""
            img = np.array(pil_image.convert("L"))  # グレースケール
            img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_LINEAR)  # 2倍拡大
            img = cv2.GaussianBlur(img, (3, 3), 0)  # ノイズ除去
            img = cv2.equalizeHist(img)  # コントラスト補正
            return Image.fromarray(img)

        def extract_text_from_docx_images(docx_path):
            """docx内の画像からOCRでテキスト抽出（前処理付き）"""
            doc = Document(docx_path)
            image_texts = []

            for rel in doc.part._rels:
                rel_obj = doc.part._rels[rel]
                if "image" in rel_obj.reltype:
                    image_data = rel_obj.target_part.blob
                    pil_image = Image.open(BytesIO(image_data)).convert("RGB")

                    # ⬇️ 画像を前処理
                    processed_image = preprocess_image_for_ocr(pil_image)

                    # PIL → PNG → bytes
                    buffered = BytesIO()
                    processed_image.save(buffered, format="PNG")
                    content = buffered.getvalue()

                    image = vision.Image(content=content)
                    response = vision_client.document_text_detection(image=image)  # ← document_text_detectionを推奨
                    texts = response.text_annotations
                    if texts:
                        image_texts.append(texts[0].description.strip())

            return "\n".join(image_texts)
        
        def extract_text_from_docx_images(docx_path):
            """docx内の画像からOCRでテキスト抽出"""
            doc = Document(docx_path)
            image_texts = []

            for rel in doc.part._rels:
                rel_obj = doc.part._rels[rel]
                if "image" in rel_obj.reltype:
                    image_data = rel_obj.target_part.blob
                    image = Image.open(BytesIO(image_data)).convert("RGB")

                    # PILイメージ → bytes for Google Vision
                    buffered = BytesIO()
                    image.save(buffered, format="PNG")
                    content = buffered.getvalue()

                    image = vision.Image(content=content)
                    #response = vision_client.text_detection(image=image)
                    response = vision_client.document_text_detection(image=image)
                    texts = response.text_annotations
                    if texts:
                        image_texts.append(texts[0].description)

            return "\n".join(image_texts)
        
        translate_llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
        def translate_to_japanese(text):
            if not text.strip():
                return text  # 空ならそのまま
            
            try:
                # 言語を検出
                detected_lang = detect(text)
                
                # 日本語の場合はそのまま返す
                if detected_lang == "ja":
                    return text
                    
                # 日本語以外の場合は翻訳
                prompt = f"""
        あなたは優秀な翻訳者です。
        以下の文章が日本語以外の言語の場合、自然な日本語に翻訳してください。
        不要な言葉(### 翻訳結果など)は除いて、翻訳結果のみ表記してください。
        ### 翻訳対象：
        {text}
        """
                result = translate_llm.invoke(prompt)
                return result.content.strip()
            except Exception as e:
                print(f"翻訳失敗: {e}")
                return text  # 失敗時は原文返す

        def process_docx_file(docx_path, llm_model=None):
            

            # === LLM設定 ===
            llm = llm_model or ChatOpenAI(model="gpt-4o-mini", temperature=0, request_timeout=30)

            # === プロンプト設定 ===
            japanese_prompt = ChatPromptTemplate.from_template("""
            以下の特許文書から、指定された項目を正確なJSON形式（Pythonの辞書形式）で抽出してください。

            ⚠️ 絶対にJSON形式のみを返答してください（文章やコメントを加えないこと）。
            ⚠️ 出力形式の例に従い、各項目の値は具体的・文書ベースで記述してください。

            出力形式の例：
            {{
                "発明の名称": "...",
                "出願人": "文章に基づいて出願人を記載してください。複数表記可（例：トヨタ自動車株式会社;株式会社豊田自動織機）",
                "発明者": "文章に基づいて発明者を記載してください。（例：特許 太郎)複数表記可（例：特許 太郎;特許 花子）",
                "公開番号": "...",
                "公開日": "文章に基づいて公開日を記載してください。記載のルールとしては西暦表示に変換して、全て半角で記載ください。(例：2025/2/6)",
                "要約": "課題、特許請求の範囲、発明が解決しようとする課題、発明の効果に関する内容を300文字以下でまとめてください。",
                "対象物": "発明の名称から特許の対象物を一言で表現してください。(例：タンク)複数表記可。（タンク;口金）",
                "対象物の最終用途": "文章に基づいて対象物が最終的に使用される用途分野（例：燃料電池車、水素インフラ）。",
                "請求項の対象": "文章に基づいて請求項の対象について一言で表現してください。（例：タンク）複数対象がある場合は複数明記ください。（例：タンク;製造方法）",
                "技術開発の背景・課題": "背景や課題を具体的に記述（文書に基づいて）",
                "技術開発課題分類": "解決しようとする技術課題を一言で表現（例：強度向上、剛性向上）。複数表記可。（強度向上;剛性向上）",
                "解決手段の概要": "課題の解決に向けての手段を具体的に記述（文書に基づいて）。",
                "解決手段分類": "解決手段を一言で表現してください（例：材料改良、形状改良、配置改良、表面加工）。複数表記可。（配置改良;表面加工）",
                "国際特許分類": "国際特許分類（例：F17C 1/06 ）があれば正確に、全て記載お願いします。複数表記可。（F16J 12/00;F17C 13/00）なければ空文字。国際特許分類は数字(△)とアルファベット(○))で構成。(○△△○ △△/△△ )。記載のルールとしては全て半角に変換して記載ください。",
                "先行技術文献"："先行技術文献の識別コードを記載してください。具体的には特許の公開番号(例：特開２０２１－１５６３１２号)などを記述してください。なければ空文字。複数表記可。（特開２０２１－１５６３１２号;特開２０２１－１１３５６０号）。"
                "Fターム": "Fタームがあれば正確に、なければ空文字",

            }}

            ## 特許全文:
            {text}
            """)
            
            chinese_prompt = ChatPromptTemplate.from_template("""
            请根据以下项目，从专利全文中提取信息，并以**日语**输出，格式为 JSON（Python 字典形式）。

            ⚠️ 只返回 JSON 格式（不要添加任何文字说明或评论）。
            ⚠️ 请按照输出示例的格式，将各项内容基于文档准确具体填写。
            ⚠️ 输出内容必须为**日语**。
            ⚠️ 请参考下方格式示例...
            

            输出格式示例：
            {{
                "发明名称": "...",
                "申请人": "请根据文章内容填写申请人，可填写多个（例如：丰田汽车株式会社;丰田自动织机株式会社）",
                "发明人": "请根据文章内容填写发明人，可填写多个（例如：特许 太郎;特许 花子）",
                "公开编号": "...",
                "公开日": "请根据文章内容填写公开日，转换为西历表示，并全部使用半角字符。（例如：2025/2/6）",
                "摘要": "请将与问题、权利要求、所要解决的问题、发明效果相关内容总结在300字以内。",
                "对象物": "请从发明名称中提取专利的对象物，用一个词表达（例如：罐体），可填写多个。（罐体;接口）",
                "对象物的最终用途": "根据文章内容，填写该对象物最终使用的用途领域（例如：燃料电池汽车、氢能基础设施）。",
                "权利要求的对象": "根据文章内容，用一个词描述权利要求对象。如有多个请全部填写（例如：罐体;制造方法）",
                "技术开发背景与课题": "请基于文档具体描述背景与所面临的课题",
                "技术开发课题分类": "将要解决的技术课题用一个词表达（例如：强度提升、刚性提升），可填写多个。（强度提升;刚性提升）",
                "解决手段概要": "请具体描述为了解决课题所采取的手段（基于文章内容）。",
                "解决手段分类": "请用一个词表达解决手段（例如：材料改良、形状改良、布局改良、表面处理），可填写多个。（布局改良;表面处理）",
                "国际专利分类": "如有国际专利分类（例如：F17C 1/06），请准确完整填写。可填写多个（F16J 12/00;F17C 13/00）。如无则为空。请转换为半角字符后填写。",
                "现有技术文献": "请填写先行技术文献的识别编号，如专利公开编号（例如：特开2021-156312号）。如无则为空。可填写多个。（特开2021-156312号;特开2021-113560号）",
                "F项分类": "如有F项分类，请准确填写。如无则为空。"
            }}

            ## 专利全文:
            {text}
            """)
            
            korean_prompt = ChatPromptTemplate.from_template("""
            다음은 특허 전문입니다. 지정된 항목에 따라 정확한 JSON 형식(Python 딕셔너리 형식)으로 정보를 추출해주세요.

            ⚠️ 반드시 JSON 형식만 반환해주세요 (문장이나 주석을 추가하지 마세요).
            ⚠️ 출력 형식 예시에 따라, 각 항목은 문서 기반으로 구체적으로 작성해주세요.

            출력 형식 예시:
            {{
                "발명의 명칭": "...",
                "출원인": "문서를 기반으로 출원인을 기재해주세요. 복수 표기 가능 (예: 도요타자동차주식회사;도요타자동직기주식회사)",
                "발명자": "문서를 기반으로 발명자를 기재해주세요. 복수 표기 가능 (예: 특허 타로;특허 하나코)",
                "공개번호": "...",
                "공개일": "문서를 기반으로 공개일을 기재해주세요. 서기로 변환하고, 전부 반각 문자로 작성해주세요. (예: 2025/2/6)",
                "요약": "과제, 특허 청구 범위, 해결하려는 문제, 발명의 효과와 관련된 내용을 300자 이내로 요약해주세요.",
                "대상물": "발명의 명칭에서 특허의 대상물을 한 단어로 표현해주세요. 복수 표기 가능 (예: 탱크;노즐)",
                "대상물의 최종 용도": "문서를 기반으로 대상물이 사용되는 최종 용도 분야를 작성해주세요 (예: 수소 인프라, 연료 전지 자동차).",
                "청구항 대상": "문서를 기반으로 청구항의 대상을 한 단어로 표현해주세요. 복수 항목이 있는 경우 모두 명시해주세요 (예: 탱크;제조 방법).",
                "기술 개발 배경 및 과제": "문서를 기반으로 배경 및 과제를 구체적으로 작성해주세요.",
                "기술 개발 과제 분류": "해결하려는 기술 과제를 한 단어로 표현해주세요 (예: 강도 향상;강성 향상). 복수 표기 가능.",
                "해결 수단 개요": "문서를 기반으로 문제 해결 수단을 구체적으로 작성해주세요.",
                "해결 수단 분류": "해결 수단을 한 단어로 표현해주세요 (예: 재료 개선, 형상 개선, 배치 개선, 표면 가공). 복수 표기 가능.",
                "국제 특허 분류": "국제 특허 분류(예: F17C 1/06)가 있으면 정확하게 모두 기재해주세요. 복수 표기 가능. 없으면 공백. 전부 반각 문자로 기입해주세요.",
                "선행 기술 문헌": "선행 기술 문헌의 식별 코드(예: 공개특허번호 등)를 기재해주세요. 없으면 공백. 복수 표기 가능.",
                "F-terms": "F-terms가 있으면 정확히, 없으면 공백."
            }}

            ## 특허 전문:
            {text}
            """)
            
            english_prompt = ChatPromptTemplate.from_template("""
            Please extract the specified items from the following patent document in accurate JSON format (Python dictionary format).

            ⚠️ Only return the JSON format (do not add any explanatory text or comments).
            ⚠️ Follow the output format example, and write each item concretely based on the content of the document.

            Output format example:
            {{
                "Title of Invention": "...",
                "Applicant": "State the applicant based on the document. Multiple entries allowed (e.g., Toyota Motor Corporation;Toyota Industries Corporation)",
                "Inventor": "State the inventor(s) based on the document. Multiple entries allowed (e.g., Tokkyo Taro;Tokkyo Hanako)",
                "Publication Number": "...",
                "Publication Date": "State the publication date based on the document. Convert to Western calendar and use half-width characters (e.g., 2025/2/6)",
                "Abstract": "Summarize the challenges, claims, problems the invention addresses, and effects of the invention in less than 300 characters.",
                "Subject Matter": "Describe the subject of the invention in one word (e.g., Tank). Multiple entries allowed (e.g., Tank;Nozzle).",
                "Final Use of the Subject": "Based on the document, describe the final application field of the subject (e.g., fuel cell vehicle, hydrogen infrastructure).",
                "Claim Target": "Describe the object of the claims in one word. If there are multiple, specify all (e.g., Tank;Manufacturing Method).",
                "Background and Issues": "Clearly describe the background and technical issues based on the document.",
                "Technical Issue Classification": "Summarize the technical issue to be solved in one word (e.g., Strength Enhancement;Rigidity Enhancement). Multiple entries allowed.",
                "Summary of Solution": "Describe in detail the means taken to solve the problem (based on the document).",
                "Solution Classification": "Express the type of solution in one word (e.g., Material Improvement, Shape Optimization, Layout Modification, Surface Treatment). Multiple entries allowed.",
                "International Patent Classification": "If available, write all International Patent Classifications accurately (e.g., F17C 1/06). Use half-width characters. Leave blank if not applicable.",
                "Prior Art Documents": "Include identifiers for prior art (e.g., publication numbers such as JP2021-156312). Leave blank if not applicable. Multiple entries allowed.",
                "F-Term": "State F-terms if available. If not, leave blank."
            }}

            ## Patent Full Text:
            {text}
            """)

            補完プロンプト = ChatPromptTemplate.from_template("""
            以下の特許全文の内、【発明の詳細な説明】、【課題】、【発明の効果】、【発明を実施するための形態】、【実施例】、【先行技術文献】をもとに、"{field}" の値を正確に日本語または数字・アルファベットで出力してください。
            ⚠️ 解説やコメントは不要です。該当する内容がない場合は空文字（""）を返してください。

            ## 特許全文:
            {text}
            """)
            
            def get_prompt_for_language(lang_code):
                if lang_code == "ja":
                    return japanese_prompt
                elif lang_code.startswith("zh"):
                    return chinese_prompt
                elif lang_code == "en":
                    return english_prompt
                elif lang_code == "ko":
                    return korean_prompt
                else:
                    return japanese_prompt  # fallback
                
            補完_chain = LLMChain(llm=llm, prompt=補完プロンプト)
            
            def normalize_keys(record, lang_code):
                key_map = {
                    "ja": {},  # 日本語ならそのまま
                    "zh": {
                        "发明名称": "発明の名称",
                        "申请人": "出願人",
                        "发明人": "発明者",
                        "公开编号": "公開番号",
                        "公开日": "公開日",
                        "摘要": "要約",
                        "对象物": "対象物",
                        "对象物的最终用途": "対象物の最終用途",
                        "权利要求的对象": "請求項の対象",
                        "技术开发背景与课题": "技術開発の背景・課題",
                        "技术开发课题分类": "技術開発課題分類",
                        "解决手段概要": "解決手段の概要",
                        "解决手段分类": "解決手段分類",
                        "国际专利分类": "国際特許分類",
                        "先行技术文献": "先行技術文献",
                        "F项分类": "Fターム",
                    },
                    "ko": {
                        "발명의 명칭": "発明の名称",
                        "출원인": "出願人",
                        "발명자": "発明者",
                        "공개번호": "公開番号",
                        "공개일": "公開日",
                        "요약": "要約",
                        "대상물": "対象物",
                        "대상물의 최종 용도": "対象物の最終用途",
                        "청구항 대상": "請求項の対象",
                        "기술 개발 배경 및 과제": "技術開発の背景・課題",
                        "기술 개발 과제 분류": "技術開発課題分類",
                        "해결 수단 개요": "解決手段の概要",
                        "해결 수단 분류": "解決手段分類",
                        "국제 특허 분류": "国際特許分類",
                        "선행 기술 문헌": "先行技術文献",
                        "F-terms": "Fターム",
                    },
                    "en": {
                        "Title of Invention": "発明の名称",
                        "Applicant": "出願人",
                        "Inventor": "発明者",
                        "Publication Number": "公開番号",
                        "Publication Date": "公開日",
                        "Abstract": "要約",
                        "Subject Matter": "対象物",
                        "Final Use of the Subject": "対象物の最終用途",
                        "Claim Target": "請求項の対象",
                        "Background and Issues": "技術開発の背景・課題",
                        "Technical Issue Classification": "技術開発課題分類",
                        "Summary of Solution": "解決手段の概要",
                        "Solution Classification": "解決手段分類",
                        "International Patent Classification": "国際特許分類",
                        "Prior Art Documents": "先行技術文献",
                        "F-Term": "Fターム",
                    }
                }

                # マッピング適用
                lang_base = lang_code.split("-")[0]  # zh-cn → zh などに対応
                mapping = key_map.get(lang_base, {})
                normalized = {}
                for k, v in record.items():
                    new_key = mapping.get(k.strip(), k.strip())  # 対応がなければそのまま
                    normalized[new_key] = v
                return normalized


            # === セクション抽出関数 ===
            def extract_sections(text, lang_code):
                if lang_code == "ja":     
                    patterns = {
                        "title": r"【発明の名称】(.+?)\n",
                        "Publication_number": r"公開番号(.+?)\n",
                        "Publication_date": r"公開日(.+?)\n",
                        "applicant": r"出願人([\s\S]+?)(?=【|$)",
                        "inventor": r"発明者([\s\S]+?)(?=【|$)",
                        "IPC": r"国際特許分類([\s\S]+?)(?=【|$)",
                        "abstract": r"【課題】([\s\S]+?)(?=【|$)",
                        "claims": r"【特許請求の範囲】([\s\S]+?)(?=【|$)",
                        "description": r"【発明の詳細な説明】([\s\S]+?)(?=【|$)",
                        "technical_problem": r"【発明が解決しようとする課題】([\s\S]+?)(?=【|$)",
                        "impact": r"【発明の効果】([\s\S]+?)(?=【|$)",
                        "prior_art": r"【先行技術文献([\s\S]+?)(?=発|$)",
                        "detail_description": r"【発明を実施するための形態】([\s\S]+?)(?=【|$)",
                        "Example": r"【実施例】([\s\S]+?)(?=【|$)",
                        "F_code": r"Fターム([\s\S]+?)(?=【|$)",
                    }
                    
                    
                elif lang_code.startswith("zh"):
                    patterns = {
                        "title": r"发明名称(.+?)\n",
                        "Publication_number": r"申请公布号 / 公开号(.+?)\n",
                        "Publication_date": r"申请公布日(.+?)\n",
                        "applicant": r"申请人([\s\S]+?)(?=【|$)",
                        "inventor": r"发明人([\s\S]+?)(?=【|$)",
                        "IPC": r"国际专利分类 (Int.Cl.)([\s\S]+?)(?=【|$)",
                        "abstract": r"背景技术([\s\S]+?)(?=【|$)",
                        "claims": r"权利要求书([\s\S]+?)(?=【|$)",
                        "description": r"说明书([\s\S]+?)(?=【|$)",
                        "technical_problem": r"具体在背景技术中表述([\s\S]+?)(?=【|$)",
                        "impact": r"发明内容内([\s\S]+?)(?=【|$)",
                        "prior_art": r"先行技术文献([\s\S]+?)(?=発|$)",
                        "detail_description": r"具体实施方式([\s\S]+?)(?=【|$)",
                        "Example": r"实施例([\s\S]+?)(?=【|$)",
                        "F_code": r"Fターム([\s\S]+?)(?=【|$)",
                    }
                    
                elif lang_code == "en":
                    patterns = {
                        "title": r"Title(.+?)\n",
                        "Publication_number": r"Publication Number(.+?)\n",
                        "Publication_date": r"Publication Date(.+?)\n",
                        "applicant": r"Applicant([\s\S]+?)(?=【|$)",
                        "inventor": r"Inventor(s)([\s\S]+?)(?=【|$)",
                        "IPC": r"International Class / Int. Cl.([\s\S]+?)(?=【|$)",
                        "abstract": r"Background of the Invention([\s\S]+?)(?=【|$)",
                        "claims": r"Claims([\s\S]+?)(?=【|$)",
                        "description": r"Detailed Description([\s\S]+?)(?=【|$)",
                        "technical_problem": r"	Background([\s\S]+?)(?=【|$)",
                        "impact": r"Background([\s\S]+?)(?=【|$)",
                        "prior_art": r"References Cited / Prior Art([\s\S]+?)(?=発|$)",
                        "detail_description": r"Detailed Description ([\s\S]+?)(?=【|$)",
                        "Example": r"Detailed Description([\s\S]+?)(?=【|$)",
                        "F_code": r"Fターム([\s\S]+?)(?=【|$)",
                    }
                
                elif lang_code == "ko":
                   patterns = {
                        "title": r"발명의\s*명칭[:：]?\s*(.+)",  # 「:」または「：」あり・なし両対応
                        "Publication_number": r"공개번호[:：]?\s*(.+)",
                        "Publication_date": r"공개일[:：]?\s*(.+)",
                        "applicant": r"출원인[:：]?\s*([\s\S]+?)(?:\n\n|$)",  # 次の空行または文末まで
                        "inventor": r"발명자[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "IPC": r"국제특허분류[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "abstract": r"배경기술[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "claims": r"(?:청구항|청구범위)[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "description": r"상세한\s*설명[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "technical_problem": r"발명의\s*목적[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "impact": r"발명의\s*효과[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "prior_art": r"선행\s*기술\s*문헌[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "detail_description": r"발명을\s*실시하기\s*위한\s*형태[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "Example": r"실시예[:：]?\s*([\s\S]+?)(?:\n\n|$)",
                        "F_code": r"F-?타?름[:：]?\s*([\s\S]+?)(?:\n\n|$)",  # Fターム (韓国特許では基本存在しないが保険で)
                    }
                    
                sections = {}
                for key, pattern in patterns.items():
                    match = re.search(pattern, text)
                    sections[key] = match.group(1).strip() if match else ""
                return sections
                    

            # === 欠損補完関数 ===
            def get_target_fields(lang_code):
                base_fields = [
                    "発明の名称", "出願人", "発明者", "公開番号", "公開日",
                    "要約", "対象物", "対象物の最終用途", "請求項の対象",
                    "技術開発の背景・課題", "技術開発課題分類", "解決手段の概要", "解決手段分類",
                    "国際特許分類", "先行技術文献"
                ]
                if lang_code == "ja":
                    base_fields.append("Fターム")
                return base_fields
            
            
            
            
            def fill_missing_fields(record, full_text):
                for field in target_fields:
                    value = record.get(field)
                    if (
                        value is None or
                        (isinstance(value, str) and not value.strip()) or
                        (isinstance(value, list) and not any(value))
                    ):
                        print(f"🔄 欠損補完中: {field}")
                        try:
                            filled = 補完_chain.run({
                                "field": field,
                                "existing_values": value,
                                "full_text": full_text
                            })
                            record[field] = filled
                        except Exception as e:
                            print(f"⚠️ {field} 補完エラー: {e}")
                return record


            # === 実行処理 ===
            filename = os.path.basename(docx_path)
            print(f"\n📘 処理中: {filename}")

            try:
                doc = Document(docx_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                try:
                    lang_code = detect(full_text)
                except:
                    lang_code = "en"
                
                main_prompt = get_prompt_for_language(lang_code)
                target_fields = get_target_fields(lang_code)
                main_chain = LLMChain(llm=llm, prompt=main_prompt)
                sections = extract_sections(full_text, lang_code)
                #merged_text = "\n".join([f"【{k}】\n{v}" for k, v in sections.items() if v])
                # ⬇️ OCRからの追加テキストを merged_text に補完
                ocr_text = extract_text_from_docx_images(docx_path)
                merged_text = full_text + "\n\n" + ocr_text  # OCRを後ろに追加
                if ocr_text:
                    merged_text += f"\n\n【OCR抽出】\n{ocr_text}"

                result = main_chain.run({"text": merged_text})
                if not result.strip().startswith("{"):
                    raise ValueError("JSON形式でない応答")

                parsed = json.loads(result)
                parsed["ファイル名"] = filename
                
                parsed = normalize_keys(parsed, lang_code)

            except Exception as e:
                print(f"⚠️ メイン抽出エラー: {e}")
                parsed = {
                    "ファイル名": filename,
                    "抽出エラー": str(e)
                }
                merged_text = full_text  # エラー時にも補完できるよう

            completed = fill_missing_fields(parsed, merged_text)
            time.sleep(1)  # レート制限対策
            return completed


        def save_to_excel(data_list, output_dir="outputs"):
            os.makedirs(output_dir, exist_ok=True)
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            filename = f"特許_抽出結果_{timestamp}_{len(data_list)}件.xlsx"
            # 全レコードに対して日本語翻訳を適用
            # ✅ 翻訳対象フィールド
            fields_to_translate = [
                "発明の名称", "要約", "対象物", "対象物の最終用途", "請求項の対象",
                "技術開発の背景・課題", "技術開発課題分類", "解決手段の概要", "解決手段分類"
            ]

            translated_data = []
            for record in data_list:
                translated_record = {}
                for k, v in record.items():
                    if isinstance(v, str) and k in fields_to_translate:
                        translated_record[k] = translate_to_japanese(v)  # 🔥 このフィールドだけ翻訳
                    else:
                        translated_record[k] = v  # その他はそのまま
                translated_data.append(translated_record)

            # 保存
            df = pd.DataFrame(translated_data)
            df.to_excel(os.path.join(output_dir, filename), index=False)
            return filename
        
        def main_app():
            if 'processed_data' not in st.session_state:
                st.session_state.processed_data = None
            if 'filename' not in st.session_state:
                st.session_state.filename = None

            uploaded_files = st.file_uploader("特許PDFをアップロードしてください（複数可）", type="pdf", accept_multiple_files=True)
            if uploaded_files and st.button("処理を開始"):
                status_area = st.empty()
                progress = st.progress(0)
                processed_data = []
                
                for i, file in enumerate(uploaded_files):
                    with open(file.name, 'wb') as f:
                        f.write(file.getbuffer())
                    status_area.info(f"{file.name} をGoogle Driveにアップロード中...")

                    pdf_id, saved_name = upload_to_drive(folder_id=PDF_input_folder, file_path=file.name)
                    progress.progress((i + 1) / len(uploaded_files) * 0.2)
                    
                    status_area.info(f"{file.name} をGoogle Docsに変換中...")
                    doc_id = convert_pdf_to_doc(pdf_id, output_folder_id=DOCS_output_folder)
                    progress.progress((i + 1) / len(uploaded_files) * 0.4)

                    docx_path = f"temp_{uuid.uuid4().hex[:8]}.docx"
                    download_doc_as_docx(doc_id, docx_path)

                    status_area.info(f"{file.name} の内容を構造化中...")
                    structured = process_docx_file(docx_path)
                    processed_data.append(structured)
                    progress.progress((i + 1) / len(uploaded_files) * 0.8)

                    os.remove(docx_path)  # 一時ファイル削除

                    status_area.info(f"{file.name} の一時ファイルを削除中...")
                    delete_from_drive(pdf_id)
                    delete_from_drive(doc_id)

                status_area.success("ダウンロードできるエクセルを生成中。もう少ししたら、ダウンロードボタンが現れます。")
                filename = save_to_excel(processed_data)
                st.session_state.processed_data = processed_data
                st.session_state.filename = filename

                progress.progress(1.0)
            
            # 1. ダウンロードボタン用のフラグを初期化
            if "downloaded" not in st.session_state:
                st.session_state.downloaded = False

            # 2. ダウンロードボタンの処理
       
            if st.session_state.processed_data and st.session_state.filename:
                with open(f"outputs/{st.session_state.filename}", "rb") as f:
                    if st.download_button("結果Excelをダウンロード", f, file_name=st.session_state.filename):
                        st.success("✅ ダウンロード完了！アップロードファイルを初期化しました。")

                        # セッション情報を削除
                        del st.session_state["processed_data"]
                        del st.session_state["filename"]

                        # 画面をクリアするために一時的な空のファイルアップローダーを表示
                        st.empty()  # 明示的な空白でUIリセット代用

        if __name__ == "__main__":
            main_app()
        

    elif mode == "データ解析システム":
        st.title("データ解析システム")
        st.write("特許要約システムで出力したエクセルファイルを用いて、データ分析・可視化を行います！")

        #st.image("./fig/データ分析システム1.jpg", use_container_width=True)
        #st.image("./fig/データ分析システム2.jpg", use_container_width=True)
        #st.image("./fig/データ分析システム3.jpg", use_container_width=True)

        
        def st_rag_langgraph():
            st.write("CSVファイルまたはxlsxファイルをアップロードしてください。")
            uploaded_file = st.file_uploader("ファイルを選択", type=["csv", "xlsx"])

            if uploaded_file is not None:
                try:
                    # ファイル拡張子をチェックして適切に処理
                    if uploaded_file.name.endswith(".csv"):
                        df = pd.read_csv(uploaded_file, index_col=0, header=0)
                
                    elif uploaded_file.name.endswith(".xlsx"):
                        df = pd.read_excel(uploaded_file, index_col=0, header=0)
                        
                    df_columns = df.columns#df_columnsに変数名を格納
                    
                    #出願日の選択
                    public_date = st.selectbox("出願日または公開日に関連する項目を選択", df_columns)
                    df[public_date] = pd.to_datetime(df[public_date], errors='coerce')
                    df["出願年"] = df[public_date].dt.year
                    df["出願件数"] = 1
                    
                    start_date = datetime(2000, 1, 1)
                    end_date = datetime(2025, 3, 31)
                    
                    term = st.slider("期間",value=[start_date,end_date])
                    
                    _df = df[(df[public_date] >= term[0]) & (df[public_date] <= term[1])]
                    
                    st.write("データプレビュー:", _df.head())
                    
                    analyze_mode = st.sidebar.radio("分析モード選択", ("特許出願件数把握", "ネットワーク分析","ワードマップ","バブルチャート"))
                
                    if analyze_mode == "特許出願件数把握":
                        applicants_list = st.selectbox("出願人に関連する項目を選択。", df_columns)
                        #データフレームのカラムを選択肢にする。複数選択
                        
                        df_clean = _df.dropna(subset=[applicants_list]).copy()
                        
                        # 出願人リスト化（セミコロン分割）
                        df_clean["applicant_list"] = df_clean[applicants_list].apply(lambda x: [a.strip() for a in str(x).split(";")])

                        # 一意な出願人名一覧（flattenして一意に）
                        flattened_applicants = [applicant for sublist in df_clean["applicant_list"] for applicant in sublist]
                        unique_values = sorted(pd.unique(flattened_applicants))

                        # 表示する出願人を選択
                        company = st.multiselect("表示する出願人", unique_values, default=unique_values )

                        # フィルター：部分一致（companyのいずれかが含まれている）
                        def company_filter(applicant_list):
                            return any(any(c in a for a in applicant_list) for c in company)

                        result = df_clean[df_clean["applicant_list"].apply(company_filter)]

                        # 共同出願判定
                        result["共同出願の有無(0:無,1:有)"] = result["applicant_list"].apply(lambda x: 1 if len(x) > 1 else 0)

                        # 集計とプロット
                        summary = result.groupby(["出願年", "共同出願の有無(0:無,1:有)"])["出願件数"].sum().reset_index()
                        fig, ax = plt.subplots(figsize=(25,12))
                        sns.barplot(data=summary, x="出願年", y="出願件数", hue="共同出願の有無(0:無,1:有)", errorbar=None)
                        ax.set_xlabel("出願年", fontsize=40)
                        ax.set_ylabel("出願件数", fontsize=40)
                        ax.tick_params(axis='y', labelsize=15)
                        ax.tick_params(axis='x', labelsize=15)
                        # 凡例のフォントサイズを変更
                        
                        legend = ax.legend(fontsize=25)  # フォントサイズを12に設定
                        legend.set_title("共同出願(無:0,有:1)", prop={'size': 25}) 

                        st.pyplot(fig)
                    
                    elif analyze_mode == "ネットワーク分析":
                                                
                        # 欠損値を削除
                        applicants_list = st.selectbox("出願人に関連する項目を選択。", df_columns)

                        # 欠損除去
                        df_clean = _df.dropna(subset=[applicants_list]).copy()
                        df_clean = df_clean.dropna(subset=["対象物"]).copy()

                        # 対象物をセミコロン分割してリスト化
                        df_clean["object_list"] = df_clean["対象物"].apply(lambda x: [a.strip() for a in str(x).split(";")])

                        # 一意な対象物をリストアップ
                        flattened_object = [obj for sublist in df_clean["object_list"] for obj in sublist]
                        unique_values = sorted(pd.unique(flattened_object))

                        # ユーザーに選ばせる
                        company = st.multiselect("ネットワークで可視化する対象物を選択してください", unique_values, default=unique_values)

                        # 部分一致フィルタ関数（対象物に1つでもcompanyのどれかが含まれていればOK）
                        def object_partial_match_filter(object_list):
                            return any(any(c in obj for obj in object_list) for c in company)

                        # フィルタ適用
                        result = df_clean[df_clean["object_list"].apply(object_partial_match_filter)]
                        
                        
                        # キーワードの分割処理
                        def split_keywords(text):
                            if pd.isna(text) or not isinstance(text, str):
                                return []
                            return [kw.strip() for kw in text.split(",") if kw.strip()]

                        result["課題リスト"] = result["技術開発課題分類"].apply(split_keywords)
                        result["技術リスト"] = result["解決手段分類"].apply(split_keywords)

                        # 共起ペア作成
                        records = []
                        for row in result.itertuples():
                            for use in row.課題リスト:
                                for tech in row.技術リスト:
                                    records.append((use, tech))

                        # ネットワーク構築
                        G = nx.Graph()
                        for use, tech in records:
                            G.add_node(use, type="課題")
                            G.add_node(tech, type="技術")
                            G.add_edge(use, tech)

                        # ノード位置計算
                        pos = nx.spring_layout(G, seed=42)

                        # Plotly向けデータ変換
                        edge_x = []
                        edge_y = []
                        for edge in G.edges():
                            x0, y0 = pos[edge[0]]
                            x1, y1 = pos[edge[1]]
                            edge_x += [x0, x1, None]
                            edge_y += [y0, y1, None]

                        node_x = []
                        node_y = []
                        node_text = []
                        node_color = []
                        for node in G.nodes():
                            x, y = pos[node]
                            node_x.append(x)
                            node_y.append(y)
                            node_text.append(node)
                            node_color.append("skyblue" if G.nodes[node]['type'] == "課題" else "lightgreen")

                        # Plotly図生成
                        edge_trace = go.Scatter(
                            x=edge_x, y=edge_y,
                            line=dict(width=0.5, color='#888'),
                            hoverinfo='none',
                            mode='lines')

                        node_trace = go.Scatter(
                            x=node_x, y=node_y,
                            mode='markers+text',
                            text=node_text,
                            textposition="top center",
                            hoverinfo='text',
                            marker=dict(
                                color=node_color,
                                size=12,
                                line_width=2))

                        fig_network = go.Figure(data=[edge_trace, node_trace],
                                    layout=go.Layout(
                                        title='課題・技術のインタラクティブネットワーク図',
                                        showlegend=False,
                                        hovermode='closest',
                                        margin=dict(b=20,l=5,r=5,t=40),
                                        xaxis=dict(showgrid=False, zeroline=False),
                                        yaxis=dict(showgrid=False, zeroline=False))
                        )

                        # ======================
                        # クラスタリング（技術要素側のみ）
                        # ======================

                        # 技術リストをベクトル化（TF-IDF）
                        tech_docs = [" ".join(techs) for techs in result["技術リスト"]]
                        vectorizer = TfidfVectorizer()
                        X = vectorizer.fit_transform(tech_docs)

                        # KMeansクラスタリング（例：5クラスタ）
                        kmeans = KMeans(n_clusters=5, random_state=42)
                        result["技術クラスタ"] = kmeans.fit_predict(X)

                        # クラスタごとの代表語（重み上位）
                        terms = vectorizer.get_feature_names_out()
                        top_keywords_per_cluster = {}
                        order_centroids = kmeans.cluster_centers_.argsort()[:, ::-1]

                        for i in range(5):
                            top_keywords = [terms[ind] for ind in order_centroids[i, :5]]
                            top_keywords_per_cluster[i] = top_keywords

                        result["技術クラスタ_代表語"] = result["技術クラスタ"].map(top_keywords_per_cluster)

                        df_clustered = result[["ファイル名", "技術リスト", "技術クラスタ", "技術クラスタ_代表語"]]
                        

                        #fig_network.show()
                        st.plotly_chart(fig_network, use_container_width=True)
                        st.write("クラスタリング結果:", df_clustered.head())
                        
                        #別ネットワーク図
                        #result = _df.dropna(subset=[applicants_list]).copy()
                        
                        #result["課題リスト"] = result["技術開発課題分類"].apply(split_keywords)
                        #result["技術リスト"] = result["解決手段分類"].apply(split_keywords)
                        result["出願人リスト"] = result["出願人"].apply(split_keywords)
                        result["発明者リスト"] = result["発明者"].apply(split_keywords)

                        # ネットワーク構築（出願人 × 用途／技術／発明者 × 技術）
                        G = nx.Graph()

                        # 出願人と用途・技術
                        for row in result.itertuples():
                            for applicant in row.出願人リスト:
                                G.add_node(applicant, type="出願人")
                                for use in row.課題リスト:
                                    G.add_node(use, type="課題")
                                    G.add_edge(applicant, use)
                                for tech in row.技術リスト:
                                    G.add_node(tech, type="技術")
                                    G.add_edge(applicant, tech)

                        # 発明者と技術要素
                        for row in result.itertuples():
                            for inventor in row.発明者リスト:
                                G.add_node(inventor, type="発明者")
                                for tech in row.技術リスト:
                                    G.add_node(tech, type="技術")
                                    G.add_edge(inventor, tech)

                        # ノード位置計算
                        pos = nx.spring_layout(G, seed=42)

                        # プロットデータ作成
                        edge_x, edge_y = [], []
                        for edge in G.edges():
                            x0, y0 = pos[edge[0]]
                            x1, y1 = pos[edge[1]]
                            edge_x += [x0, x1, None]
                            edge_y += [y0, y1, None]

                        node_x, node_y, node_text, node_color = [], [], [], []
                        for node in G.nodes():
                            x, y = pos[node]
                            node_x.append(x)
                            node_y.append(y)
                            node_text.append(node)
                            node_type = G.nodes[node]['type']
                            if node_type == "出願人":
                                node_color.append("lightcoral")
                            elif node_type == "課題":
                                node_color.append("skyblue")
                            elif node_type == "技術":
                                node_color.append("lightgreen")
                            elif node_type == "発明者":
                                node_color.append("gold")

                        # Plotlyネットワーク描画
                        edge_trace = go.Scatter(
                            x=edge_x, y=edge_y,
                            line=dict(width=0.5, color='#888'),
                            hoverinfo='none',
                            mode='lines')

                        node_trace = go.Scatter(
                            x=node_x, y=node_y,
                            mode='markers+text',
                            text=node_text,
                            textposition="top center",
                            hoverinfo='text',
                            marker=dict(
                                color=node_color,
                                size=12,
                                line_width=2))

                        fig_nt = go.Figure(data=[edge_trace, node_trace],
                                    layout=go.Layout(
                                        title='出願人・発明者 × 技術要素・課題のネットワーク図',
                                        showlegend=False,
                                        hovermode='closest',
                                        margin=dict(b=20,l=5,r=5,t=40),
                                        xaxis=dict(showgrid=False, zeroline=False),
                                        yaxis=dict(showgrid=False, zeroline=False))
                        )
                        st.plotly_chart(fig_nt, use_container_width=True)
                        
                        #別ネットワーク図
                        #result = _df.dropna(subset=[applicants_list]).copy()
                                        
                        #result["技術リスト"] = result["解決手段分類"].apply(split_keywords)
                        #result["出願人リスト"] = result["出願人"].apply(split_keywords)
                        #result["発明者リスト"] = result["発明者"].apply(split_keywords)

                        # TF-IDF + KMeansクラスタリング（技術要素ベース）
                        tech_docs = [" ".join(techs) for techs in result["技術リスト"]]
                        vectorizer = TfidfVectorizer()
                        X = vectorizer.fit_transform(tech_docs)
                        kmeans = KMeans(n_clusters=5, random_state=42)
                        result["技術クラスタ"] = kmeans.fit_predict(X)

                        # ==========================
                        # 出願人・発明者ごとのクラスタ分布集計
                        # ==========================

                        # 出願人 × 技術クラスタ
                        applicant_cluster_records = []
                        for row in result.itertuples():
                            for applicant in row.出願人リスト:
                                applicant_cluster_records.append((applicant, row.技術クラスタ))
                        df_applicant_cluster = pd.DataFrame(applicant_cluster_records, columns=["出願人", "技術クラスタ"])
                        df_applicant_cluster_summary = df_applicant_cluster.value_counts().unstack(fill_value=0)

                        # 発明者 × 技術クラスタ
                        inventor_cluster_records = []
                        for row in result.itertuples():
                            for inventor in row.発明者リスト:
                                inventor_cluster_records.append((inventor, row.技術クラスタ))
                        df_inventor_cluster = pd.DataFrame(inventor_cluster_records, columns=["発明者", "技術クラスタ"])
                        df_inventor_cluster_summary = df_inventor_cluster.value_counts().unstack(fill_value=0)

                        # ==========================
                        # エッジ重み付きネットワーク図（出願人 × 技術要素）
                        # ==========================

                        # 出願人と技術の共起数をカウント
                        edge_counter = {}
                        for row in result.itertuples():
                            for applicant in row.出願人リスト:
                                for tech in row.技術リスト:
                                    pair = (applicant, tech)
                                    edge_counter[pair] = edge_counter.get(pair, 0) + 1

                        # グラフ構築（重み付き）
                        G = nx.Graph()
                        for (applicant, tech), weight in edge_counter.items():
                            G.add_node(applicant, type="出願人")
                            G.add_node(tech, type="技術")
                            G.add_edge(applicant, tech, weight=weight)

                        # ノード位置
                        pos = nx.spring_layout(G, seed=42)

                        # エッジ線（重みで太さ調整）
                        edge_x, edge_y, edge_width = [], [], []
                        for u, v, data in G.edges(data=True):
                            x0, y0 = pos[u]
                            x1, y1 = pos[v]
                            edge_x += [x0, x1, None]
                            edge_y += [y0, y1, None]
                            edge_width.append(data['weight'])

                        edge_trace = go.Scatter(
                            x=edge_x, y=edge_y,
                            line=dict(width=1, color='gray'),
                            hoverinfo='none',
                            mode='lines'
                        )

                        # ノード描画
                        node_x, node_y, node_text, node_color = [], [], [], []
                        for node in G.nodes():
                            x, y = pos[node]
                            node_x.append(x)
                            node_y.append(y)
                            node_text.append(node)
                            node_type = G.nodes[node]['type']
                            node_color.append("lightcoral" if node_type == "出願人" else "lightgreen")

                        node_trace = go.Scatter(
                            x=node_x, y=node_y,
                            mode='markers+text',
                            text=node_text,
                            textposition="top center",
                            hoverinfo='text',
                            marker=dict(color=node_color, size=12, line_width=2)
                        )

                        # 図作成
                        fig_weighted = go.Figure(data=[edge_trace, node_trace],
                            layout=go.Layout(
                                title='出願人 × 技術要素の重み付きネットワーク図',
                                showlegend=False,
                                hovermode='closest',
                                margin=dict(b=20,l=5,r=5,t=40),
                                xaxis=dict(showgrid=False, zeroline=False),
                                yaxis=dict(showgrid=False, zeroline=False))
                        )
                        
                        st.plotly_chart(fig_weighted, use_container_width=True)
                        
                    elif analyze_mode == "ワードマップ":
                        st.write("対象物")    
                        # データから発明等の名称を取得
                        appnames = _df['対象物'].values

                        # Janomeトークナイザを初期化
                        tokenizer = Tokenizer()
                        words = []

                        # 名詞の連続を検出して複合名詞として結合
                        for appname in appnames:
                            tokens = tokenizer.tokenize(appname)
                            noun_phrase = []  # 複合名詞用のリスト
                            for token in tokens:
                                if token.part_of_speech.startswith('名詞'):
                                    noun_phrase.append(token.surface)  # 名詞を追加
                                else:
                                    if noun_phrase:  # 名詞が連続していた場合、結合してリストに追加
                                        words.append("".join(noun_phrase))
                                        noun_phrase = []  # リセット
                            # 最後に残った名詞を追加
                            if noun_phrase:
                                words.append("".join(noun_phrase))

                        # 単語の出現頻度を計算
                        df_words = pd.Series(words).value_counts()
                        word_counts = df_words.to_dict()

                        # 単語の頻度を棒グラフで表示
                        fig,ax = plt.subplots(figsize=(20, 13))
                        head_20 = df_words.iloc[:20].copy()
                        ax.barh(y=head_20.index, width=head_20.values, color="orange")
                                            
                        # グラフのタイトルとラベル
                        ax.set_title("頻出単語のトップ20", fontsize=40)
                        ax.set_xlabel("頻度", fontsize=40)
                        ax.set_ylabel("単語", fontsize=40)

                        # y軸のラベルを調整
                        #ax.set_yticks(head_20.index)
                        #ax.set_yticklabels(head_20.index, rotation=0, fontsize=30)
                        ax.tick_params(axis='y', labelsize=30)
                        ax.tick_params(axis='x', labelsize=30)
                        
                        st.pyplot(fig)
                        
                        # 日本語フォントを指定
                        font_path1 = "./font/NotoSansJP-Regular.ttf"  # 適切なパスを指定

                        # ワードクラウドを生成
                        wordcloud = WordCloud(
                            background_color='white', 
                            width=800, 
                            height=600, 
                            font_path=font_path1
                        )

                        wordcloud.generate_from_frequencies(word_counts)

                        # ワードクラウドを表示
                        plt.figure(figsize=(10, 8))
                        plt.imshow(wordcloud, interpolation='bilinear')
                        plt.axis('off')

                        # Streamlitで表示
                        st.pyplot(plt)
                        st.write("請求項の対象")
                        appnames = _df['請求項の対象'].values

                        # Janomeトークナイザを初期化
                        tokenizer = Tokenizer()
                        words = []

                        # 名詞の連続を検出して複合名詞として結合
                        for appname in appnames:
                            tokens = tokenizer.tokenize(appname)
                            noun_phrase = []  # 複合名詞用のリスト
                            for token in tokens:
                                if token.part_of_speech.startswith('名詞'):
                                    noun_phrase.append(token.surface)  # 名詞を追加
                                else:
                                    if noun_phrase:  # 名詞が連続していた場合、結合してリストに追加
                                        words.append("".join(noun_phrase))
                                        noun_phrase = []  # リセット
                            # 最後に残った名詞を追加
                            if noun_phrase:
                                words.append("".join(noun_phrase))

                        # 単語の出現頻度を計算
                        df_words = pd.Series(words).value_counts()
                        word_counts = df_words.to_dict()

                        # 単語の頻度を棒グラフで表示
                        fig,ax = plt.subplots(figsize=(20, 13))
                        head_20 = df_words.iloc[:20].copy()
                        ax.barh(y=head_20.index, width=head_20.values, color="orange")
                                            
                        # グラフのタイトルとラベル
                        ax.set_title("頻出単語のトップ20", fontsize=40)
                        ax.set_xlabel("頻度", fontsize=40)
                        ax.set_ylabel("単語", fontsize=40)

                        # y軸のラベルを調整
                        #ax.set_yticks(head_20.index)
                        #ax.set_yticklabels(head_20.index, rotation=0, fontsize=30)
                        ax.tick_params(axis='y', labelsize=30)
                        ax.tick_params(axis='x', labelsize=30)
                        
                        st.pyplot(fig)
                        


                        # ワードクラウドを生成
                        wordcloud = WordCloud(
                            background_color='white', 
                            width=800, 
                            height=600, 
                            font_path=font_path1
                        )

                        wordcloud.generate_from_frequencies(word_counts)

                        # ワードクラウドを表示
                        plt.figure(figsize=(10, 8))
                        plt.imshow(wordcloud, interpolation='bilinear')
                        plt.axis('off')
                        
                        

                        # Streamlitで表示
                        st.pyplot(plt)
                        st.write("技術開発課題分類")
                        appnames = _df['技術開発課題分類'].values

                        # Janomeトークナイザを初期化
                        tokenizer = Tokenizer()
                        words = []

                        # 名詞の連続を検出して複合名詞として結合
                        for appname in appnames:
                            tokens = tokenizer.tokenize(appname)
                            noun_phrase = []  # 複合名詞用のリスト
                            for token in tokens:
                                if token.part_of_speech.startswith('名詞'):
                                    noun_phrase.append(token.surface)  # 名詞を追加
                                else:
                                    if noun_phrase:  # 名詞が連続していた場合、結合してリストに追加
                                        words.append("".join(noun_phrase))
                                        noun_phrase = []  # リセット
                            # 最後に残った名詞を追加
                            if noun_phrase:
                                words.append("".join(noun_phrase))

                        # 単語の出現頻度を計算
                        df_words = pd.Series(words).value_counts()
                        word_counts = df_words.to_dict()

                        # 単語の頻度を棒グラフで表示
                        fig,ax = plt.subplots(figsize=(20, 13))
                        head_20 = df_words.iloc[:20].copy()
                        ax.barh(y=head_20.index, width=head_20.values, color="orange")
                                            
                        # グラフのタイトルとラベル
                        ax.set_title("頻出単語のトップ20", fontsize=40)
                        ax.set_xlabel("頻度", fontsize=40)
                        ax.set_ylabel("単語", fontsize=40)

                        # y軸のラベルを調整
                        #ax.set_yticks(head_20.index)
                        #ax.set_yticklabels(head_20.index, rotation=0, fontsize=30)
                        ax.tick_params(axis='y', labelsize=30)
                        ax.tick_params(axis='x', labelsize=30)
                        
                        st.pyplot(fig)
                        


                        # ワードクラウドを生成
                        wordcloud = WordCloud(
                            background_color='white', 
                            width=800, 
                            height=600, 
                            font_path=font_path1
                        )

                        wordcloud.generate_from_frequencies(word_counts)

                        # ワードクラウドを表示
                        plt.figure(figsize=(10, 8))
                        plt.imshow(wordcloud, interpolation='bilinear')
                        plt.axis('off')
                        
                        

                        # Streamlitで表示
                        st.pyplot(plt)
                        
                        st.write("解決手段分類")
                        appnames = _df['解決手段分類'].values

                        # Janomeトークナイザを初期化
                        tokenizer = Tokenizer()
                        words = []

                        # 名詞の連続を検出して複合名詞として結合
                        for appname in appnames:
                            tokens = tokenizer.tokenize(appname)
                            noun_phrase = []  # 複合名詞用のリスト
                            for token in tokens:
                                if token.part_of_speech.startswith('名詞'):
                                    noun_phrase.append(token.surface)  # 名詞を追加
                                else:
                                    if noun_phrase:  # 名詞が連続していた場合、結合してリストに追加
                                        words.append("".join(noun_phrase))
                                        noun_phrase = []  # リセット
                            # 最後に残った名詞を追加
                            if noun_phrase:
                                words.append("".join(noun_phrase))

                        # 単語の出現頻度を計算
                        df_words = pd.Series(words).value_counts()
                        word_counts = df_words.to_dict()

                        # 単語の頻度を棒グラフで表示
                        fig,ax = plt.subplots(figsize=(20, 13))
                        head_20 = df_words.iloc[:20].copy()
                        ax.barh(y=head_20.index, width=head_20.values, color="orange")
                                            
                        # グラフのタイトルとラベル
                        ax.set_title("頻出単語のトップ20", fontsize=40)
                        ax.set_xlabel("頻度", fontsize=40)
                        ax.set_ylabel("単語", fontsize=40)

                        # y軸のラベルを調整
                        #ax.set_yticks(head_20.index)
                        #ax.set_yticklabels(head_20.index, rotation=0, fontsize=30)
                        ax.tick_params(axis='y', labelsize=30)
                        ax.tick_params(axis='x', labelsize=30)
                        
                        st.pyplot(fig)
                        


                        # ワードクラウドを生成
                        wordcloud = WordCloud(
                            background_color='white', 
                            width=800, 
                            height=600, 
                            font_path=font_path1
                        )

                        wordcloud.generate_from_frequencies(word_counts)

                        # ワードクラウドを表示
                        plt.figure(figsize=(10, 8))
                        plt.imshow(wordcloud, interpolation='bilinear')
                        plt.axis('off')
                        
                        

                        # Streamlitで表示
                        st.pyplot(plt)

                    elif analyze_mode == "バブルチャート":
                         # 欠損値を削除
                        applicants_list = st.selectbox("出願人に関連する項目を選択。", df_columns)

                        # 欠損除去
                        df_clean = _df.dropna(subset=[applicants_list]).copy()
                        df_clean = df_clean.dropna(subset=["対象物"]).copy()

                        # 対象物をセミコロン分割してリスト化
                        df_clean["object_list"] = df_clean["対象物"].apply(lambda x: [a.strip() for a in str(x).split(";")])

                        # 一意な対象物をリストアップ
                        flattened_object = [obj for sublist in df_clean["object_list"] for obj in sublist]
                        unique_values = sorted(pd.unique(flattened_object))

                        # ユーザーに選ばせる
                        company = st.multiselect("ネットワークで可視化する対象物を選択してください", unique_values, default=unique_values)

                        # 部分一致フィルタ関数（対象物に1つでもcompanyのどれかが含まれていればOK）
                        def object_partial_match_filter(object_list):
                            return any(any(c in obj for obj in object_list) for c in company)

                        # フィルタ適用
                        result = df_clean[df_clean["object_list"].apply(object_partial_match_filter)]
                        
                        
                        # キーワードの分割処理
                        def split_keywords(text):
                            if pd.isna(text) or not isinstance(text, str):
                                return []
                            return [kw.strip() for kw in text.split(",") if kw.strip()]

                        result["課題リスト"] = result["技術開発課題分類"].apply(split_keywords)
                        result["技術リスト"] = result["解決手段分類"].apply(split_keywords)
                        
                        bubble_data = []
                        for _, row in result.iterrows():
                            company_name = row[applicants_list]
                            for use in row.課題リスト:
                                for tech in row.技術リスト:
                                    bubble_data.append({
                                        '技術': tech,
                                        '課題': use,
                                        '出願人': company_name
                                    })

                        # データフレームに変換して集計
                        bubble_df = pd.DataFrame(bubble_data)
                        bubble_count = bubble_df.groupby(['技術', '課題']).size().reset_index(name='出願件数')

                        # 上位の技術と課題を取得（あまりに多いと見づらくなるため）
                        top_techs = bubble_df['技術'].value_counts().nlargest(15).index.tolist()
                        top_issues = bubble_df['課題'].value_counts().nlargest(15).index.tolist()

                        # フィルタリング
                        filtered_bubble = bubble_count[
                            bubble_count['技術'].isin(top_techs) & 
                            bubble_count['課題'].isin(top_issues)
                        ]

                        # バブルチャート作成
                        fig_bubble = go.Figure()

                        fig_bubble.add_trace(go.Scatter(
                            x=filtered_bubble['技術'],
                            y=filtered_bubble['課題'],
                            mode='markers',
                            marker=dict(
                                size=filtered_bubble['出願件数'] * 10,  # サイズは適宜調整
                                sizemode='area',
                                sizeref=2. * max(filtered_bubble['出願件数']) / (40.**2),  # バブルサイズ調整
                                sizemin=4,
                                color=filtered_bubble['出願件数'],
                                colorscale='Viridis',
                                showscale=True,
                                colorbar=dict(title='出願件数')
                            ),
                            text=[f'技術: {tech}<br>課題: {use}<br>出願数: {count}' 
                                for tech, use, count in zip(filtered_bubble['技術'], filtered_bubble['課題'], filtered_bubble['出願件数'])],
                            hoverinfo='text'
                        ))

                        fig_bubble.update_layout(
                            title='技術-課題のバブルチャート（バブルサイズ：出願件数）',
                            xaxis=dict(
                                title='技術',
                                categoryorder='total ascending'
                            ),
                            yaxis=dict(
                                title='課題',
                                categoryorder='total ascending'
                            ),
                            height=800,
                            width=900
                        )

                        # プロット表示
                        st.plotly_chart(fig_bubble, use_container_width=True)

                        # 出願人別の集計も見せる
                        st.subheader("出願人別の集計")
                        company_counts = bubble_df['出願人'].value_counts().reset_index()
                        company_counts.columns = ['出願人', '出願件数']
                        st.dataframe(company_counts)

                
                except Exception as e:
                    st.error(f"ファイルを読み込む際にエラーが発生しました: {e}")
                    
        if __name__ == "__main__":
            st_rag_langgraph()
    
    elif mode == "対比表作成システム":
        

        # --- Drive関連関数 ---
        def upload_to_drive(folder_id, file_path, mime_type='application/pdf'):
            unique_name = f"{os.path.splitext(os.path.basename(file_path))[0]}_{int(time.time())}_{uuid.uuid4().hex[:6]}.pdf"
            media = MediaFileUpload(file_path, mimetype=mime_type)
            file_metadata = {'name': unique_name, 'parents': [folder_id]}
            uploaded = drive_service.files().create(body=file_metadata, media_body=media, fields='id,name').execute()
            return uploaded['id'], uploaded['name']

        def delete_from_drive(file_id):
            try:
                drive_service.files().delete(fileId=file_id).execute()
            except HttpError as e:
                print(f"削除失敗: {e}")

        def convert_pdf_to_doc(file_id, output_folder_id):
            copied = drive_service.files().copy(fileId=file_id, body={
                'mimeType': 'application/vnd.google-apps.document',
                'parents': [output_folder_id]
            }).execute()
            return copied['id']

        def download_doc_as_docx(file_id, local_path):
            try:
                request = drive_service.files().export_media(fileId=file_id,
                                                            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                with open(local_path, 'wb') as f:
                    f.write(request.execute())
                return local_path
            except HttpError as error:
                print(f"ドキュメントのダウンロード失敗: {error}")
                return None

        # --- OCR前処理 ---
        def preprocess_image_for_ocr(pil_image):
            img = np.array(pil_image.convert("L"))
            img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_LINEAR)
            img = cv2.GaussianBlur(img, (3, 3), 0)
            img = cv2.equalizeHist(img)
            return Image.fromarray(img)

        # --- OCR抽出関数 ---
        def extract_text_from_docx_images(docx_path):
            doc = Document(docx_path)
            image_texts = []
            for rel in doc.part._rels:
                rel_obj = doc.part._rels[rel]
                if "image" in rel_obj.reltype:
                    image_data = rel_obj.target_part.blob
                    pil_image = Image.open(BytesIO(image_data)).convert("RGB")
                    buffered = BytesIO()
                    pil_image.save(buffered, format="PNG")
                    content = buffered.getvalue()

                    image = vision.Image(content=content)
                    response = vision_client.document_text_detection(image=image)
                    texts = response.text_annotations
                    if texts:
                        image_texts.append(texts[0].description.strip())
            return "\n".join(image_texts)

        # --- チャンク分割 ---
        def split_text_into_chunks(text, chunk_size=1000, overlap=100):
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=overlap,
                separators=["\n\n", "\n", ".", "。", "、", " "]
            )
            return splitter.split_text(text)

        # --- FAISSインデックス作成 ---
        def create_faiss_index(text_chunks, embedding_model):
            docs = [LangChainDocument(page_content=chunk) for chunk in text_chunks]
            return FAISS.from_documents(docs, embedding_model)

        # --- 請求項を意味単位で分割 ---
        def split_claims_into_chunks(claims_list, llm):
            split_results = []
            for claim in claims_list:
                prompt = ChatPromptTemplate.from_template("""
        あなたは特許請求項を自然な意味のまとまりごとに分割する専門家です。
        以下の請求項を100〜200文字程度で自然に分割してください。

        【請求項】:
        {claim}

        【出力形式】:
        - 分割1: ...
        - 分割2: ...
        """)
                chain = prompt | llm
                try:
                    response = chain.invoke({"claim": claim})
                    split_text = response.content.strip()
                    split_texts = []
                    for line in split_text.split("\n"):
                        if "- 分割" in line:
                            _, text = line.split(":", 1)
                            split_texts.append(text.strip())
                    split_results.append(split_texts)
                except Exception as e:
                    print(f"分割エラー: {e}")
                    split_results.append([claim])
            return split_results
        
        def main():
            # --- Streamlitアプリ本体 ---
            
            st.title("特許対比表作成ツール（FAISS＋スコア＋共通点整理版）")
            st.image("./fig/対比表作成システム1.jpg", use_container_width=True)
            st.image("./fig/対比表作成システム2.jpg", use_container_width=True)
            st.image("./fig/対比表作成システム3.jpg", use_container_width=True)

            uploaded_excel = st.file_uploader("請求項リスト（Excel）をアップロード", type=["xlsx"])
            uploaded_pdfs = st.file_uploader("比較対象の特許PDFをアップロード（複数可）", type=["pdf"], accept_multiple_files=True)


            if uploaded_excel and uploaded_pdfs:
                claims_df = pd.read_excel(uploaded_excel)
                claim_column = st.selectbox("請求項が記載されている列を選んでください", claims_df.columns)

                if st.button("✨ 対比表作成スタート！"):
                    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.0)
                    embedding_model = OpenAIEmbeddings(model="text-embedding-ada-002")
                    progress = st.progress(0.1)

                    # --- ステップ1: 請求項を分割 ---
                    claim_texts = claims_df[claim_column].tolist()
                    split_claims = split_claims_into_chunks(claim_texts, llm)

                    split_rows = []
                    for idx, splits in enumerate(split_claims):
                        for split_idx, split_text in enumerate(splits, 1):
                            split_rows.append({
                                "元請求項No.": idx + 1,
                                "元請求項本文": claim_texts[idx],
                                "分割No.": split_idx,
                                "分割文": split_text,
                            })
                    split_df = pd.DataFrame(split_rows)

                    # --- ステップ2: PDFごとに処理 ---
                    for pdf_file in uploaded_pdfs:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                            tmp.write(pdf_file.read())
                            tmp_path = tmp.name

                        # ファイルアップロード
                        file_id, file_name = upload_to_drive(GOOGLE_DRIVE_FOLDER_ID, tmp_path)

                        # PDF→Google Docs変換
                        doc_id = convert_pdf_to_doc(file_id, GOOGLE_DRIVE_FOLDER_ID)

                        # DOCXダウンロード
                        docx_local_path = download_doc_as_docx(doc_id, tmp_path.replace(".pdf", ".docx"))
                        
                        text = extract_text_from_docx_images(docx_local_path)

                        delete_from_drive(file_id)
                        delete_from_drive(doc_id)
                        os.remove(tmp_path)
                        os.remove(docx_local_path)

                        if not text.strip():
                            st.warning(f"⚠️ ファイル {pdf_file.name} はOCR結果が空でした。スキップします。")
                            continue

                        chunks = split_text_into_chunks(text, chunk_size=1000)
                        if not chunks:
                            st.warning(f"⚠️ ファイル {pdf_file.name} はチャンク分割できませんでした。スキップします。")
                            continue

                        faiss_index = create_faiss_index(chunks, embedding_model)

                        # 3列セット作成
                        result_col_text = f"{pdf_file.name}_類似原文"
                        result_col_score = f"{pdf_file.name}_スコア"
                        result_col_summary = f"{pdf_file.name}_共通点整理"

                        split_df[result_col_text] = ""
                        split_df[result_col_score] = ""
                        split_df[result_col_summary] = ""

                        for idx, row in split_df.iterrows():
                            query = row["分割文"]
                            results = faiss_index.similarity_search_with_score(query, k=1)

                            if results:
                                best_doc, best_score = results[0]
                                best_text = best_doc.page_content

                                split_df.at[idx, result_col_text] = best_text
                                split_df.at[idx, result_col_score] = round(best_score, 3)

                                prompt = ChatPromptTemplate.from_template("""
            以下の請求項分割文と、抽出された特許本文を比較し、
            - 技術分野
            - 材料
            - 工程
            - 特徴的な条件
            を丁寧に列挙してください。

            【請求項分割文】:
            {split_text}

            【抽出された特許本文】:
            {best_text}

            【出力形式】:
            - 技術分野: ...
            - 材料: ...
            - 工程: ...
            - 条件: ...
            """)
                                chain = prompt | llm

                                try:
                                    response = chain.invoke({"split_text": query, "best_text": best_text})
                                    split_df.at[idx, result_col_summary] = response.content.strip()
                                except Exception as e:
                                    split_df.at[idx, result_col_summary] = f"エラー: {e}"

                            else:
                                split_df.at[idx, result_col_text] = "類似なし"
                                split_df.at[idx, result_col_score] = ""
                                split_df.at[idx, result_col_summary] = "整理対象なし"

                        st.progress((uploaded_pdfs.index(pdf_file)+1) / len(uploaded_pdfs), text=f"{uploaded_pdfs.index(pdf_file)+1}/{len(uploaded_pdfs)}ファイル目 処理中")
                        
                    progress.progress(0.8)
                    # --- ステップ3: 出力
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        claims_df.to_excel(writer, index=False, sheet_name="原文")
                        split_df.to_excel(writer, index=False, sheet_name="分割後＋対比結果")

                    st.success("✅ 完了しました！")
                    progress.progress(1.0)
                    st.download_button(
                        label="📥 エクセルファイルをダウンロード",
                        data=output.getvalue(),
                        file_name="請求項対比結果_FAISS_3列版.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

        if __name__ == "__main__":
            main()


